#!/usr/bin/python3
# coding=utf-8

import glob, os, fnmatch, operator, re, argparse

from docx import Document
from docx.shared import Cm, Pt
from PIL import Image
from pathlib import Path
from datetime import date

version = "1.0"

path = "./"
properties = []
result = []

resultFile = ""
templateFile = ""
reportFile = ""

caseno = ""
client = ""
worker = ""

#Order of columns in report, columns must be specified to appear in report
order = ['Path', 'Created', 'Modified', 'Accessed', 'Author', 'Create Time', 'Created (Meta-data)', 'Modified (Meta-data)', 'Producer', 'Last Printed', 'Last Saved By', 'Last Saved Time', 'Revision Number', 'Total Editing Time', 'Submit Time', 'Subject', 'Attachment Count', 'From', 'To', 'CC', 'BCC', 'Unread', 'Unsent', 'L-Size (bytes)', 'MD5', 'Deleted', 'Carved', 'Item #', ]

removePath = 4 #Remove part of part before n:th /

key_translate = {  "Path":   "Path", 
          "Created":   "Created",
          "Modified":  "Modified",
          "Accessed":  "Accessed",
          "Author":  "Author",
          "Create Time":  "Create time",
          "Last Printed":  "Last printed by",
          "Last Saved By":  "Last saved by",
          "Last Saved Time":  "Last saved time",
          "Revision Number":  "Revision",
          "Total Editing Time":  "Total editing time",
          "Submit Time":  "Submit time",
          "Subject":  "Subject",
          "Attachment Count": "Attchment count",
          "From": "From",
          "To": "To",
          "CC": "CC",
          "BCC": "BCC",
          "Unread": "Unread",
          "Unsent": "Unsent",
          "L-Size (bytes)": "Size (bytes)",
          "MD5": "Hash (MD5)",
          "Created (Meta-data)": "Document created",
          "Modified (Meta-data)": "Document modified",
          "Producer": "Tool",
          "Deleted": "Deleted",
          "Carved": "Carved",
          "Item #": "Item # (FTK)"
          }
          
value_translate = {  "True":    "Ja",
          "true":    "Ja",
          "TRUE":    "Ja",
          "False":  "Nej",
          "false":  "Nej",
          "FALSE":  "Nej"
          }

exif_translate = {  "Author": "Author",
          "CreateDate": "Created (Meta-data)",
          "ModifyDate": "Modified (Meta-data)",
          "Producer": "Producer"
}

#Class: Item
#Represent one file/document with multiple properties/metadata
class Item(dict):
  __slots__ = properties
  
  def __init__(self, values):
    i = 0
    for key in properties:
      if len(key) > 0:
        key = key.strip()
        self[key] = values[i].strip()
      i += 1
   
  #Sort
  def __lt__(self, other):
    return self['Path'] < other['Path']
      
#getItemNumber, return itemnumber from path
def getItemNumber(path:str) -> int:

  returnValue = 0

  if(("[" in path) and ("]" in path)):
    returnValue = int(path.rsplit("[", 1)[1].split("]")[0])

  return returnValue

#getResultFromFile, read result from file
def getResultFromFile(textfile):

  global properties

  returnValue = []
  
  with open(textfile, "r") as f:
    for line in f.readlines():
      if not properties: #header in first row
        #text columns
        properties = line.replace("\r\n","").replace("\n","").split("\t")
      else:
        #itemnumber = line.split("\t")[0]
        #path = line.split("\t")[1]
        returnValue.append(Item(line.replace("\r\n", "").split("\t")))

  #Add exif columns
  for item in returnValue:
    for value in exif_translate.values():
      if value not in item.keys():
        item[value] = ""

  return returnValue

#getItemPosition, return position of item if exists otherwise -1
def getItemPosition(itemnumber:int) -> int:
  returnValue = -1
  
  for i, item in enumerate(result):
    if(int(item["Item #"]) == itemnumber):
      returnValue = i
      break 

  return returnValue

#getResultFromExif, extend result with exif information
def getResultFromExif(exiffolder):

  for exif_file in find("*.exif.txt", exiffolder):
    path = str(Path(os.path.basename(exif_file)))
    itemnumber = re.compile(r"(\[)([0-9]{1,8})(\])").search(path).group(2) #...[itemnumber]...exif.txt
    
    i = getItemPosition(int(itemnumber))

    if i > -1:
      item = result[i]
      
      with open(exif_file, "r") as f:
        for line in f.readlines():
          key = line.split(": ")[0]
          value = line.split(": ")[1].strip()

          if key in exif_translate.keys():
            #Don't overwrite stored values
            if len(item[_translate(key, exif_translate)]) == 0:
              #Replace : with - in dates
              if re.compile(r"[0-9]{4}:[0-9]{2}:[0-9]{2} ").search(value):
                item[_translate(key, exif_translate)] = value.split(" ")[0].replace(":", "-") + " " + value.split(" ")[1]
              else:
                item[_translate(key, exif_translate)] = value

#getResult, get result for given itemnumber
def getResult(itemnumber:int) -> Item:

  returnValue = ""

  for item in result:
    if(int(item["Item #"]) == itemnumber):
      returnValue = item
      
  return returnValue
  
def find(pattern, path):
  returnValue = []
  
  for root, _, files in os.walk(path):
    for name in files:
      if(fnmatch.fnmatch(name, pattern)):
        returnValue.append(os.path.join(root, name))
        
  return returnValue

#getThumbnail, get thumbnail path for given itemnumber
def getThumbnail(itemnumber, path):
  return find("*[[]" + str(itemnumber) + "[]].*jp*g", path)

#getThumbnailPath, return path to thumbnail based on given path extension
def getThumbnailPath(path):
  returnValue = ""
     
  file = os.path.basename(path)
  ext = os.path.splitext(file)[1]
  filename = os.path.splitext(file)[0]
  
  type = ext.upper()[1:]
  
  if(path.lower().endswith(".emlx.msg.htm")):
    type = "EMLX"
  
  if(path.lower().endswith(".eml.htm")):
    type = "EML"

  if(path.lower().endswith(".html")):
    type = "HTML"
    
  if(path.lower().endswith(".htm")):
    type = "HTM"
    
  if(path.lower().endswith(".lnk.htm")):
    type = "LNK"
 
  if(type in ["DOC", "DOCX", "ODS", "ODT", "XLS", "XLSX", "XLT", "TXT"]):
    returnValue = path.replace(file, "Converted/Img/" + filename + ".pdf.jpg")
  elif(type in ["EMLX", "EML", "LNK", "HTML", "HTM"]):
    returnValue = path.replace(file, "Converted/Img/" + file + ".pdf.jpg")
  elif(type in ["PDF", "TIF", "TIFF", "PNG"]):
    returnValue = path.replace(file, "Converted/Img/" + file + ".jpg")
  elif(type in ["JPEG", "JPG"]):
    returnValue = path.replace(file, "Converted/Img/" + file)
  
  return returnValue

#getImageWidth, return scaled image width for thumbnail
def getImageWidth(height, rows, headerChars):

  returnValue = 1100
    
  returnValue -= (rows + 3)*25 #many rows #0.1
  
  if(headerChars > 40):
    returnValue -= (headerChars/40 - 1) * 25 #long paths 40 0.1
  
  if(returnValue > 700):
    returnValue = 600
  
  if(returnValue > height): #use original height if calculated height is smaller
    returnValue = height

  #print("return cm=" + str(Cm(returnValue * (2.54/150))) + ", return=" + str(returnValue) + ", height=" + str(height) + ", rows=" + str(rows) + ", chars=" + str(headerChars))
  
  return Cm(returnValue * (2.54/150))
  
#_translate, lookup key in table and return value if key exists otherwise key
def _translate(key, table):
  
  returnValue = key

  if(key in table):
    returnValue = table[key]

  return returnValue
  
#getKey, return value from key_translate table if key exists  
def getKey(key):
  return _translate(key, key_translate)

#getValue, return value from value_translate table if value exists
def getValue(value):
  return _translate(value, value_translate)

#report, generate report
def report(items, output_file):

  document = Document(templateFile)
  keys = [] #List of used keys to remove description for not used keys

  #Replace dynamic text in header
  for section in document.sections:
    header = section.header
    for table in header.tables:
      for row in table.rows:
        for cell in row.cells:
          for paragraph in cell.paragraphs:
            if any(x in paragraph.text for x in ["{date}", "{caseno}"]):
              paragraph.text = paragraph.text.replace("{date}", str(date.today())).replace("{caseno}", str(caseno))

  #Replace dynamic text in document
  for table in document.tables:
    for row in table.rows:
      for cell in row.cells:
        for paragraph in cell.paragraphs:
          paragraph.text = paragraph.text.replace("{worker}", str(worker)).replace("{client}", str(client))

  #Create one page with metadata + thumbnail for every item
  for item in items:
    print(".", end="", flush=True)
    table = document.add_table(0, 2)
    rowCount = 0
    
    for key in order: #sorted(item):
      if(len(item[key]) > 0 and item[key] not in ["False", "n/a"]):
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(getKey(str(key))) + ":"
        
        if str(getKey(str(key))) not in keys:
          keys.append(str(getKey(str(key))))
        
        if(key == "Path"):
          cells[1].text = str('/'.join(map(str, item[key].split("/")[removePath:]))) #remove first part of path (image, volume name, ...)
        elif(key == "Total Editing Time"):
          cells[1].text = item[key].replace("hours", "timmar").replace("minutes", "minuter").replace("seconds", "sekunder")
        else:
          cells[1].text = str(getValue(item[key]))
    
        rowCount += 1

    table.style = "DR_FilEgenskaper"
    
    #Set table width
    widths = (Cm(6), Cm(13))
    for row in table.rows:
      for idx, width in enumerate(widths):
        row.cells[idx].width = width
  
    #Add thumbnail
    img_path = item["ThumbnailPath"]
    
    if os.path.isfile(img_path):
      img = Image.open(img_path)
      document.add_paragraph()
      
      table = document.add_table(0,1)
      row = table.add_row()
      cells = row.cells
      cells[0].text = str("Thumbnail (first page)")
      table.style = "DR_FilEgenskaper"
    
      for row in table.rows:
        row.cells[0].width = Cm(19)
      
      try:
        paragraph = document.add_paragraph("") #Empty space between header and image
        document.add_picture(str(img_path), width=getImageWidth(img.size[1], rowCount, len(item['Path'])))
        
      except Exception as ex:
        print("Error document.add_picture", img_path, ex)
    
    if(item != items[-1]): #Last page = no page break
      document.add_page_break()
  
  #Remove property descriptions which is not used
  for table in document.tables:
    if table.rows[0].cells[0].paragraphs[0].text == "Egenskap": #Property description table found
      for row in table.rows:
        if row.cells[0].paragraphs[0].text not in keys:
          row._element.getparent().remove(row._element)

  document.save(output_file)

if __name__ == "__main__":

  parser = argparse.ArgumentParser(prog="createReport.py", description="Author:\tDaniel Rådesjö (daniel.radesjo@gmail.com)\n\thttps://github.com/daniel-radesjo", formatter_class=argparse.RawTextHelpFormatter)
  parser.add_argument("-r", "--result", action="store", type=str, help="textfile with metadata", default="result.txt")
  parser.add_argument("-t", "--template", action="store", type=str, help="report template", default="template.docx")
  parser.add_argument("-o", "--report", action="store", type=str, help="output file", default="report.docx")
  parser.add_argument("-n", "--case", action="store", type=str, help="case number in report", default="")
  parser.add_argument("-c", "--client", action="store", type=str, help="client in report", default="")
  parser.add_argument("-w", "--worker", action="store", type=str, help="worker in report", default="")
  parser.add_argument("-v", "--version", action="version", version="%(prog)s " + version)

  args = parser.parse_args()

  resultFile = args.result
  templateFile = args.template
  reportFile = args.report
  caseno = args.case
  client = args.client
  worker = args.worker

  print("createReport.py " + version)
  print("-" * 80)
  for key,value in vars(args).items():
    print(key + ": " + value)
  print("-" * 80)

  result = getResultFromFile(resultFile)
  getResultFromExif("exif")
  items = []

  for dir in sorted(glob.glob("./*/")):
    for path in sorted(glob.glob(dir + "*")):

      filename = os.path.basename(path)

      if(os.path.isfile(path) and os.path.isfile(getThumbnailPath(path))) and "/exif/" not in path:
        item = getResult(getItemNumber(path))
        item["ThumbnailPath"] = getThumbnailPath(path)       
        items.append(item)

  items.sort()
  print("Creating report...", end="", flush=True)
  report(items, reportFile)  
  print("\nReport created")
